from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


TODO_TEXT = '\u5f85\u8865\u5145\u3002'


def _apply_document_fonts(document: Document) -> None:
    for style_name in ['Normal', 'Heading 1', 'Heading 2']:
        style = document.styles[style_name]
        style.font.name = 'Times New Roman'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        if style_name == 'Normal':
            style.font.size = Pt(11)
        elif style_name == 'Heading 1':
            style.font.size = Pt(16)
        elif style_name == 'Heading 2':
            style.font.size = Pt(13)


def build_experiment_doc(output_path: Path, title: str, sections: list[str]) -> Path:
    document = Document()
    _apply_document_fonts(document)
    document.add_heading(title, level=1)
    for section in sections:
        document.add_heading(section, level=2)
        document.add_paragraph(TODO_TEXT)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def build_checkpoint_doc(output_path: Path, title: str, sections: list[tuple[str, list[str]]]) -> Path:
    document = Document()
    _apply_document_fonts(document)
    document.add_heading(title, level=1)
    for heading, paragraphs in sections:
        document.add_heading(heading, level=2)
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path
