from pathlib import Path
import subprocess
import sys

from docx import Document
import pytest

from battery_thesis.documents import build_experiment_doc
from battery_thesis.plot_style import apply_academic_plot_style


PROJECT_ROOT = Path(__file__).resolve().parents[1]
PYTHON = sys.executable


def test_apply_academic_plot_style_sets_minus_fix_and_has_font_fallbacks():
    style = apply_academic_plot_style()

    assert style['axes.unicode_minus'] is False
    assert style['font.family']
    assert 'DejaVu Sans' in style['font.family']


def test_build_experiment_doc_creates_expected_headings(tmp_path: Path):
    output_path = tmp_path / 'experiment_summary.docx'

    build_experiment_doc(
        output_path=output_path,
        title='实验摘要',
        sections=['实验设计', '结果分析'],
    )

    assert output_path.exists()

    document = Document(output_path)
    texts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    assert '实验摘要' in texts
    assert '实验设计' in texts
    assert '结果分析' in texts
    assert '待补充。' in texts


def test_plotting_script_runs_as_subprocess(tmp_path: Path):
    input_path = tmp_path / 'plot_data.csv'
    png_path = tmp_path / 'plot.png'
    svg_path = tmp_path / 'plot.svg'
    input_path.write_text('category,series,value\nsd,f1,0.8\nsd,recall,0.7\n', encoding='utf-8')

    result = subprocess.run(
        [
            PYTHON,
            'scripts/plotting/plot_recognition_bar.py',
            '--input',
            str(input_path),
            '--output_png',
            str(png_path),
            '--output_svg',
            str(svg_path),
        ],
        cwd=PROJECT_ROOT,
        capture_output=True,
        text=True,
        check=False,
    )

    assert result.returncode == 0, result.stderr
    assert png_path.exists()
    assert svg_path.exists()


def test_deliverable_docx_files_do_not_contain_garbled_text():
    deliverables = Path('deliverables')
    if not deliverables.exists():
        pytest.skip('requires generated deliverable DOCX files, which are not committed')
    for name in ['experiment_summary.docx', 'implementation_checkpoint.docx', 'innovation_summary.docx', 'method_summary.docx']:
        document = Document(deliverables / name)
        texts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        assert texts
        assert not any('?' in text or '\ufffd' in text or '??' in text for text in texts)
